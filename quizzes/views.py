from __future__ import annotations

import logging
from io import BytesIO

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpRequest, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.crypto import get_random_string
from django.views import View

from materials.models import Material

from classrooms.models import QuizAssignment
from .models import Quiz, QuizAttempt, QuizQuestion, QuizShareLink
from .services import GeminiError, generate_quiz
from .utils import grade_quiz_submission

logger = logging.getLogger(__name__)


class GenerateQuizView(LoginRequiredMixin, View):
    def post(self, request, material_id: int):
        material = get_object_or_404(Material, pk=material_id, owner=request.user)
        try:
            question_count = int(request.POST.get("question_count", 10))
        except (TypeError, ValueError):
            question_count = 10
        question_count = max(1, min(question_count, 50))
        question_type = request.POST.get("question_type", QuizQuestion.QuestionType.MULTIPLE_CHOICE)
        difficulty = request.POST.get("difficulty", Quiz.Difficulty.MEDIUM)
        timer_minutes = request.POST.get("timer_minutes", "").strip()
        
        # Validate difficulty
        if difficulty not in [Quiz.Difficulty.EASY, Quiz.Difficulty.MEDIUM, Quiz.Difficulty.HARD]:
            difficulty = Quiz.Difficulty.MEDIUM
        
        # Parse timer_minutes (optional)
        timer_value = None
        if timer_minutes:
            try:
                timer_value = int(timer_minutes)
                if timer_value < 0:
                    timer_value = None
            except (ValueError, TypeError):
                timer_value = None

        # Normalize and validate question type values from the form (also accept a few aliases)
        def _normalize_qtype(raw: str | None) -> str:
            if not raw:
                return QuizQuestion.QuestionType.MULTIPLE_CHOICE
            val = str(raw).strip().lower()
            if val in (QuizQuestion.QuestionType.MULTIPLE_CHOICE, 'multiple', 'mc', 'multiple_choice'):
                return QuizQuestion.QuestionType.MULTIPLE_CHOICE
            if val in (QuizQuestion.QuestionType.TRUE_FALSE, 'true/false', 'truefalse', 'tf', 'boolean'):
                return QuizQuestion.QuestionType.TRUE_FALSE
            if val in (QuizQuestion.QuestionType.FILL_IN_BLANK, 'fill-in-blank', 'fill in the blank', 'fib'):
                return QuizQuestion.QuestionType.FILL_IN_BLANK
            return QuizQuestion.QuestionType.MULTIPLE_CHOICE

        # Log the raw values received from the form
        logger.info(f"Form data - question_count: {request.POST.get('question_count')}, question_type: {request.POST.get('question_type')}")

        question_type = _normalize_qtype(question_type)
            
        quiz = Quiz.objects.create(
            owner=request.user,
            material=material,
            status=Quiz.Status.PROCESSING,
            difficulty=difficulty,
            timer_minutes=timer_value,
            settings={
                "question_count": question_count,
                "question_type": question_type,
                "difficulty": difficulty,
                "timer_minutes": timer_value,
            },
        )
        redirect_url = reverse('materials:upload') + '#queue'
        try:
            # Log the values being passed to generate_quiz
            logger.info(f"Calling generate_quiz with question_count={question_count}, question_type={question_type}, difficulty={difficulty}")
            payload = generate_quiz(
                material,
                question_count=question_count,
                question_type=question_type,
                difficulty=difficulty
            )
            reduced_from = payload.pop('reduced_from', None) if isinstance(payload, dict) else None
        except GeminiError as exc:
            quiz.status = Quiz.Status.ERROR
            quiz.error_message = str(exc)
            quiz.save(update_fields=["status", "error_message", "updated_at"])
            messages.error(request, f"Quiz generation failed: {exc}")
        else:
            questions = payload.get("questions", [])
            if reduced_from:
                actual = payload.get('question_count', len(questions))
                messages.warning(
                    request,
                    f'Gemini could only generate {actual} questions instead of {reduced_from}.',
                )
            if not questions:
                quiz.status = Quiz.Status.ERROR
                quiz.error_message = 'Gemini did not return any questions.'
                quiz.save(update_fields=['status', 'error_message', 'updated_at'])
                messages.error(request, 'Gemini did not return any questions.')
            else:
                quiz.title = payload.get('quiz_title', material.title or 'Generated Quiz')
                quiz.model_name = getattr(settings, 'GEMINI_MODEL', '')
                quiz.question_count = payload.get('question_count', len(questions))
                quiz.status = Quiz.Status.READY
                quiz.save(update_fields=['title', 'model_name', 'question_count', 'status', 'updated_at'])
                for index, item in enumerate(questions):
                    choices = item.get("choices", [])
                    correct_answer = ""

                    # Always use the user's selected question type, not what Gemini returns
                    # This ensures the question type matches what the user requested
                    question_type_for_item = question_type

                    if question_type_for_item == QuizQuestion.QuestionType.MULTIPLE_CHOICE:
                        correct_index = item.get("correct_index", 0)
                        if choices:
                            position = min(max(correct_index, 0), max(len(choices) - 1, 0))
                            correct_answer = choices[position]
                    elif question_type_for_item == QuizQuestion.QuestionType.TRUE_FALSE:
                        correct_answer = str(item.get("correct_answer", False)).lower()
                    else:  # fill_in_blank
                        correct_answer = item.get("correct_answer", "")

                    QuizQuestion.objects.create(
                        quiz=quiz,
                        prompt=item.get("prompt", ""),
                        choices=choices,
                        correct_answer=correct_answer,
                        explanation=item.get("explanation", ""),
                        order=index,
                        question_type=question_type_for_item,
                    )
                messages.success(request, "Quiz generated successfully.")
                redirect_url = reverse('quizzes:detail', args=[quiz.pk])
        return redirect(redirect_url)


class QuizListView(LoginRequiredMixin, View):
    template_name = 'quizzes/list.html'

    def get(self, request, material_id: int | None = None):
        from django.db.models import Q
        
        quizzes = Quiz.objects.filter(owner=request.user).select_related('material')
        material = None
        if material_id:
            quizzes = quizzes.filter(material_id=material_id)
            material = get_object_or_404(Material, pk=material_id, owner=request.user)
        
        # Search functionality
        search_query = request.GET.get('search', '')
        if search_query:
            quizzes = quizzes.filter(
                Q(title__icontains=search_query) |
                Q(material__title__icontains=search_query)
            )
        
        # Filter by status
        status_filter = request.GET.get('status', '')
        if status_filter:
            quizzes = quizzes.filter(status=status_filter)
        
        # Filter by material
        material_filter = request.GET.get('material', '')
        if material_filter:
            quizzes = quizzes.filter(material_id=material_filter)
        
        quizzes = quizzes.order_by('-created_at')
        
        # Get available materials for filter dropdown
        available_materials = Material.objects.filter(owner=request.user).order_by('-created_at')
        
        return render(request, self.template_name, {
            'quizzes': quizzes,
            'material': material,
            'search_query': search_query,
            'status_filter': status_filter,
            'material_filter': material_filter,
            'available_materials': available_materials,
            'status_choices': Quiz.Status.choices,
        })


class QuizScheduleView(LoginRequiredMixin, View):
    """
    Calendar-like schedule of quizzes with classroom due dates for the current user.
    """

    template_name = "quizzes/schedule.html"

    def get(self, request: HttpRequest) -> HttpResponse:
        from django.utils import timezone

        now = timezone.now()
        is_instructor = getattr(request.user, "is_instructor", lambda: False)()

        # Base assignments queryset: quizzes tied to classes the user is involved with.
        assignments_qs = (
            QuizAssignment.objects.select_related("quiz", "classroom")
            .filter(quiz__status=Quiz.Status.READY)
        )

        # Students see assignments for classes they are members of.
        # Instructors see assignments for classes they own.
        if is_instructor:
            assignments_qs = assignments_qs.filter(classroom__owner=request.user)
        else:
            assignments_qs = assignments_qs.filter(classroom__memberships__user=request.user)

        assignments_qs = assignments_qs.distinct().order_by("due_at", "created_at")
        assignments = list(assignments_qs)

        # Preload attempts so we can highlight submitted quizzes without N+1 queries.
        attempts_by_assignment: dict[int, QuizAttempt] = {
            attempt.assignment_id: attempt
            for attempt in QuizAttempt.objects.filter(
                user=request.user,
                assignment__in=assignments_qs,
            )
        }

        # Attach a lightweight "user_attempt" attribute to each assignment instance
        # for easy access in templates.
        for assignment in assignments:
            setattr(assignment, "user_attempt", attempts_by_assignment.get(assignment.id))

        upcoming: list[QuizAssignment] = []
        no_due: list[QuizAssignment] = []
        past: list[QuizAssignment] = []

        for assignment in assignments:
            due_at = assignment.due_at
            user_attempt = getattr(assignment, "user_attempt", None)

            if is_instructor:
                # Instructors: only show assignments that are not past due.
                if due_at and due_at < now:
                    continue
                if due_at is None:
                    no_due.append(assignment)
                else:
                    upcoming.append(assignment)
            else:
                # Students: hide assignments once they have at least one attempt.
                if user_attempt is not None:
                    continue
                if due_at is None:
                    no_due.append(assignment)
                elif due_at >= now:
                    upcoming.append(assignment)
                else:
                    past.append(assignment)

        context = {
            "upcoming_assignments": upcoming,
            "no_due_assignments": no_due,
            "past_assignments": past,
            "now": now,
        }
        return render(request, self.template_name, context)


class QuizDetailView(LoginRequiredMixin, View):
    template_name = 'quizzes/detail.html'

    def get_quiz(self, request, pk: int) -> Quiz:
        return get_object_or_404(
            Quiz.objects.prefetch_related('questions').select_related('material'),
            pk=pk,
            owner=request.user,
        )

    def get(self, request, pk: int):
        quiz = self.get_quiz(request, pk)
        questions = list(quiz.questions.all())
        entries = [{'question': question, 'result': None} for question in questions]
        share_link = None
        if getattr(request.user, "is_instructor", lambda: False)():
            share_link = quiz.share_links.filter(is_active=True).order_by("-created_at").first()
        share_url = reverse("quizzes:shared_take", args=[share_link.token]) if share_link else None
        return render(request, self.template_name, {
            'quiz': quiz,
            'questions': questions,
            'entries': entries,
            'share_link': share_link,
            'share_url': request.build_absolute_uri(share_url) if share_url else None,
        })

    def post(self, request, pk: int):
        quiz = self.get_quiz(request, pk)
        questions = list(quiz.questions.all())
        entries, score, total, percent = grade_quiz_submission(questions, request.POST)
        share_link = None
        if getattr(request.user, "is_instructor", lambda: False)():
            share_link = quiz.share_links.filter(is_active=True).order_by("-created_at").first()
        share_url = reverse("quizzes:shared_take", args=[share_link.token]) if share_link else None
        return render(request, self.template_name, {
            'quiz': quiz,
            'questions': questions,
            'entries': entries,
            'submitted': True,
            'score': score,
            'total': total,
            'percent': percent,
            'share_link': share_link,
            'share_url': request.build_absolute_uri(share_url) if share_url else None,
        })


class QuizShareLinkView(LoginRequiredMixin, View):
    """Create, regenerate, or disable a direct share link for a quiz."""

    def post(self, request, pk: int):
        quiz = get_object_or_404(Quiz, pk=pk, owner=request.user)
        action = request.POST.get("action", "create")
        active_link = quiz.share_links.filter(is_active=True).order_by("-created_at").first()

        if action == "deactivate":
            if active_link:
                active_link.is_active = False
                active_link.save(update_fields=["is_active"])
                messages.info(request, "Share link disabled.")
            else:
                messages.info(request, "No active share link to disable.")
        elif action == "regenerate":
            if active_link:
                active_link.token = get_random_string(24)
                active_link.is_active = True
                active_link.save(update_fields=["token", "is_active"])
                messages.success(request, "Share link regenerated.")
            else:
                QuizShareLink.objects.create(quiz=quiz, created_by=request.user)
                messages.success(request, "Share link created.")
        else:
            if not active_link:
                QuizShareLink.objects.create(quiz=quiz, created_by=request.user)
                messages.success(request, "Share link created.")
            else:
                messages.info(request, "Share link already active.")

        return redirect("quizzes:detail", pk=quiz.pk)


class SharedQuizTakeView(LoginRequiredMixin, View):
    """Allow students to take a quiz via a direct share link."""
    template_name = "quizzes/shared_take.html"

    def get(self, request, token: str):
        share_link = get_object_or_404(
            QuizShareLink.objects.select_related("quiz", "quiz__material", "created_by"),
            token=token,
            is_active=True,
        )
        if share_link.quiz.status != Quiz.Status.READY:
            messages.error(request, "This quiz is not ready yet.")
            return redirect("quizzes:list")
        questions = list(share_link.quiz.questions.all())
        entries = [{'question': question, 'result': None} for question in questions]

        # If the quiz is assigned to a class the user is in, enforce the deadline.
        assignment = (
            QuizAssignment.objects.filter(
                quiz=share_link.quiz,
                classroom__memberships__user=request.user,
            )
            .order_by("due_at")
            .first()
        )
        now = timezone.now()
        deadline_seconds = None
        if assignment and assignment.due_at:
            remaining = (assignment.due_at - now).total_seconds()
            deadline_seconds = int(remaining) if remaining > 0 else 0

        time_limit_seconds = None
        if share_link.quiz.timer_minutes:
            time_limit_seconds = share_link.quiz.timer_minutes * 60
        if deadline_seconds is not None:
            time_limit_seconds = (
                min(time_limit_seconds, deadline_seconds) if time_limit_seconds is not None else deadline_seconds
            )

        existing_attempt = QuizAttempt.objects.filter(share_link=share_link, user=request.user).first()
        submitted = False
        score = total = percent = None
        if existing_attempt:
            submitted = True
            score = existing_attempt.score
            total = existing_attempt.total_questions
            percent = float(existing_attempt.percent)
            for entry in entries:
                qid = str(entry["question"].id)
                stored = existing_attempt.answers.get(qid, {})
                entry["result"] = {
                    "user_answer": stored.get("user_answer", ""),
                    "correct": stored.get("correct", False),
                }

        allow_submit = share_link.quiz.status == Quiz.Status.READY and not submitted
        if assignment and assignment.due_at and now >= assignment.due_at and not submitted:
            allow_submit = False
            messages.error(request, "This shared quiz is past the classroom deadline.")
        return render(request, self.template_name, {
            "share_link": share_link,
            "quiz": share_link.quiz,
            "questions": questions,
            "entries": entries,
            "submitted": submitted,
            "score": score,
            "total": total,
            "percent": percent,
            "allow_submit": allow_submit,
            "deadline_seconds": deadline_seconds,
            "time_limit_seconds": time_limit_seconds,
        })

    def post(self, request, token: str):
        share_link = get_object_or_404(
            QuizShareLink.objects.select_related("quiz", "quiz__material", "created_by"),
            token=token,
            is_active=True,
        )
        if share_link.quiz.status != Quiz.Status.READY:
            messages.error(request, "This quiz is not ready yet.")
            return redirect("quizzes:list")

        assignment = (
            QuizAssignment.objects.filter(
                quiz=share_link.quiz,
                classroom__memberships__user=request.user,
            )
            .order_by("due_at")
            .first()
        )
        now = timezone.now()
        if assignment and assignment.due_at and now >= assignment.due_at and not request.POST.get("auto_submit"):
            messages.error(request, "This shared quiz can no longer be submitted; the classroom deadline has passed.")
            return redirect("quizzes:list")

        existing_attempt = QuizAttempt.objects.filter(share_link=share_link, user=request.user).first()
        if existing_attempt:
            messages.info(request, "You have already submitted this quiz.")
            questions = list(share_link.quiz.questions.all())
            entries = []
            for question in questions:
                qid = str(question.id)
                stored = existing_attempt.answers.get(qid, {})
                entries.append(
                    {
                        "question": question,
                        "result": {
                            "user_answer": stored.get("user_answer", ""),
                            "correct": stored.get("correct", False),
                        },
                    }
                )
            return render(request, self.template_name, {
                "share_link": share_link,
                "quiz": share_link.quiz,
                "questions": questions,
                "entries": entries,
                "submitted": True,
                "score": existing_attempt.score,
                "total": existing_attempt.total_questions,
                "percent": float(existing_attempt.percent),
                "allow_submit": False,
                "deadline_seconds": None,
                "time_limit_seconds": None,
            })

        questions = list(share_link.quiz.questions.all())
        entries, score, total, percent = grade_quiz_submission(questions, request.POST)
        answers_payload = {
            str(entry["question"].id): {
                "user_answer": entry["result"]["user_answer"],
                "correct": entry["result"]["correct"],
            }
            for entry in entries
        }

        QuizAttempt.objects.update_or_create(
            share_link=share_link,
            user=request.user,
            defaults={
                "quiz": share_link.quiz,
                "score": score,
                "total_questions": total,
                "percent": percent,
                "answers": answers_payload,
            },
        )

        messages.success(request, "Quiz submitted.")
        return render(request, self.template_name, {
            "share_link": share_link,
            "quiz": share_link.quiz,
            "questions": questions,
            "entries": entries,
            "submitted": True,
            "score": score,
            "total": total,
            "percent": percent,
            "allow_submit": False,
            "deadline_seconds": None,
            "time_limit_seconds": None,
        })


class QuizExportPDFView(LoginRequiredMixin, View):
    """Export quiz as PDF."""
    def get(self, request, pk: int):
        quiz = get_object_or_404(Quiz, pk=pk, owner=request.user)
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#4F46E5',
                spaceAfter=12,
            )
            story.append(Paragraph(quiz.title or "Quiz", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Quiz info
            info_text = f"Material: {quiz.material.title}<br/>"
            info_text += f"Difficulty: {quiz.get_difficulty_display()}<br/>"
            info_text += f"Questions: {quiz.question_count}"
            story.append(Paragraph(info_text, styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Questions (without answers - for student use)
            for idx, question in enumerate(quiz.questions.all(), 1):
                story.append(Paragraph(f"<b>Question {idx}:</b> {question.prompt}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
                
                if question.question_type == 'multiple_choice' and question.choices:
                    # Show choices with letters (A, B, C, D, etc.)
                    choice_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
                    for i, choice in enumerate(question.choices):
                        if i < len(choice_letters):
                            story.append(Paragraph(f"  {choice_letters[i]}. {choice}", styles['Normal']))
                elif question.question_type == 'true_false':
                    story.append(Paragraph("  A. True", styles['Normal']))
                    story.append(Paragraph("  B. False", styles['Normal']))
                else:
                    # Fill in the blank - show blank line
                    story.append(Paragraph("  Answer: _______________________", styles['Normal']))
                
                story.append(Spacer(1, 0.3*inch))  # Space for student to write answer
            
            doc.build(story)
            buffer.seek(0)
            
            response = HttpResponse(buffer.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="quiz_{quiz.pk}.pdf"'
            return response
        except ImportError:
            messages.error(request, "PDF export requires 'reportlab' package. Please install it.")
            return redirect('quizzes:detail', pk=pk)


class QuizExportDOCXView(LoginRequiredMixin, View):
    """Export quiz as DOCX."""
    def get(self, request, pk: int):
        quiz = get_object_or_404(Quiz, pk=pk, owner=request.user)
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(quiz.title or "Quiz", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title.runs[0].font.color.rgb = RGBColor(79, 70, 229)
            
            # Quiz info
            doc.add_paragraph(f"Material: {quiz.material.title}")
            doc.add_paragraph(f"Difficulty: {quiz.get_difficulty_display()}")
            doc.add_paragraph(f"Questions: {quiz.question_count}")
            doc.add_paragraph()  # Blank line
            
            # Questions (without answers - for student use)
            for idx, question in enumerate(quiz.questions.all(), 1):
                doc.add_heading(f"Question {idx}", level=2)
                doc.add_paragraph(question.prompt)
                
                if question.question_type == 'multiple_choice' and question.choices:
                    # Show choices with letters (A, B, C, D, etc.)
                    choice_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
                    for i, choice in enumerate(question.choices):
                        if i < len(choice_letters):
                            doc.add_paragraph(f"{choice_letters[i]}. {choice}")
                elif question.question_type == 'true_false':
                    doc.add_paragraph("A. True")
                    doc.add_paragraph("B. False")
                else:
                    # Fill in the blank - show blank line
                    doc.add_paragraph("Answer: _______________________")
                
                doc.add_paragraph()  # Blank line for student to write answer
                doc.add_paragraph()  # Extra spacing
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="quiz_{quiz.pk}.docx"'
            return response
        except ImportError:
            messages.error(request, "DOCX export requires 'python-docx' package. Please install it.")
            return redirect('quizzes:detail', pk=pk)


class QuizDeleteView(LoginRequiredMixin, View):
    """
    Delete a quiz owned by the current user.
    Used from the dashboard and quizzes list screens.
    """

    def post(self, request: HttpRequest, pk: int) -> HttpResponse:
        quiz = get_object_or_404(Quiz, pk=pk, owner=request.user)
        quiz.delete()
        messages.success(request, "Quiz deleted.")

        next_url = request.POST.get("next")
        if next_url:
            return redirect(next_url)
        return redirect("quizzes:list")
